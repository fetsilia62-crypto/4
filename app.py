from __future__ import annotations

import os
import re
import shutil
import sys
import tempfile
import threading
from datetime import date
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

import fitz
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from docx import Document


APP_TITLE = "Паспорт → договор ГПХ"
PLACEHOLDER = "\u2002" * 5

FIELDS = [
    ("contract_number", "Номер договора"),
    ("contract_date", "Дата договора"),
    ("fio", "ФИО"),
    ("passport_series_number", "Паспорт: серия и номер"),
    ("issued_by", "Кем выдан"),
    ("issue_date", "Дата выдачи"),
    ("department_code", "Код подразделения"),
    ("birth_date", "Дата рождения"),
    ("registration_address", "Адрес регистрации"),
    ("phone", "Телефон"),
    ("inn", "ИНН"),
    ("snils", "СНИЛС"),
    ("bank_name", "Банк"),
    ("correspondent_account", "Корреспондентский счёт"),
    ("settlement_account", "Расчётный счёт"),
    ("bik", "БИК"),
]

ISSUER_WORDS = (
    "МВД", "УФМС", "ОТДЕЛ", "ОТДЕЛОМ", "ОТДЕЛЕНИЕМ",
    "РОССИИ", "ГУ", "УПРАВЛЕНИ", "ПЕТЕРБУРГ", "МОСКВ",
)

NAME_STOPWORDS = {
    "РОССИЙСКАЯ", "ФЕДЕРАЦИЯ", "ПАСПОРТ", "МВД", "УФМС", "РОССИИ",
    "ОТДЕЛА", "ОТДЕЛОМ", "ОТДЕЛЕНИЕМ", "УПРАВЛЕНИЕ", "УПРАВЛЕНИЯ",
    "ОБЛАСТИ", "ОБЛ", "РАЙОНЕ", "РАЙОН", "САНКТ", "ПЕТЕРБУРГ",
    "ПЕТЕРБУРГУ", "МОСКВА", "МОСКВЕ", "ДАТА", "ВЫДАЧИ", "КОД",
    "ПОДРАЗДЕЛЕНИЯ", "ЛИЧНЫЙ", "ПОДПИСЬ", "МУЖ", "ЖЕН", "ПОЛ",
    "МЕСТО", "РОЖДЕНИЯ", "ГОР", "ГОРОД", "РЕСП", "КРАЙ", "ЧЕЛ",
}


def app_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


def configure_tesseract() -> tuple[Path, Path | None]:
    base = app_dir()
    bundled = base / "tesseract" / "tesseract.exe"
    if bundled.exists():
        exe = bundled
    else:
        found = shutil.which("tesseract")
        if not found:
            raise RuntimeError(
                "Не найден модуль OCR Tesseract.\n"
                "Запускай программу из всей папки PassportContract, а не выноси один EXE отдельно."
            )
        exe = Path(found)

    pytesseract.pytesseract.tesseract_cmd = str(exe)

    # В собранной Windows-версии языковые файлы лежат рядом с Tesseract.
    # При запуске исходников допускаем системную папку tessdata.
    bundled_tessdata = exe.parent / "tessdata"
    if (bundled_tessdata / "rus.traineddata").exists():
        return exe, bundled_tessdata

    env_tessdata = os.environ.get("TESSDATA_PREFIX")
    candidates = [
        Path(env_tessdata) if env_tessdata else None,
        Path("/usr/share/tesseract-ocr/5/tessdata"),
        Path("/usr/share/tessdata"),
    ]
    for candidate in candidates:
        if candidate and (candidate / "rus.traineddata").exists():
            return exe, candidate

    # Системный Tesseract может сам знать путь к языковым файлам.
    try:
        if "rus" in pytesseract.get_languages(config=""):
            return exe, None
    except Exception:
        pass

    raise RuntimeError("Не найден русский языковой файл OCR rus.traineddata.")


def render_pdf(pdf_path: Path, out_dir: Path) -> list[Path]:
    images: list[Path] = []
    with fitz.open(str(pdf_path)) as pdf:
        for index, page in enumerate(pdf):
            pix = page.get_pixmap(matrix=fitz.Matrix(3, 3), alpha=False)
            image_path = out_dir / f"page_{index + 1}.png"
            pix.save(str(image_path))
            images.append(image_path)
    return images


def prepare_image(image: Image.Image) -> Image.Image:
    image = ImageOps.exif_transpose(image).convert("RGB")
    if image.width < 1800:
        factor = 1800 / image.width
        image = image.resize(
            (int(image.width * factor), int(image.height * factor)),
            Image.Resampling.LANCZOS,
        )
    gray = ImageOps.grayscale(image)
    gray = ImageOps.autocontrast(gray, cutoff=1)
    gray = ImageEnhance.Contrast(gray).enhance(1.35)
    gray = gray.filter(ImageFilter.SHARPEN)
    return gray


def text_score(text: str) -> float:
    upper = text.upper()
    keywords = (
        "РОССИЙСК", "ФЕДЕРАЦ", "ПАСПОРТ", "УФМС", "МВД",
        "ДАТА", "ВЫДАЧ", "PNRUS", "ПОДРАЗДЕЛ",
    )
    keyword_score = sum(80 for word in keywords if word in upper)
    cyrillic = len(re.findall(r"[А-ЯЁа-яё]", text))
    digits = len(re.findall(r"\d", text))
    return keyword_score + cyrillic + digits * 0.5


def ocr_image(image_path: Path, tessdata: Path | None) -> str:
    source = prepare_image(Image.open(image_path))
    best_text = ""
    best_score = -1.0
    config = "--oem 1 --psm 6"
    if tessdata is not None:
        config = f'--tessdata-dir "{tessdata}" ' + config

    for angle in (0, 90, 270, 180):
        candidate = source if angle == 0 else source.rotate(angle, expand=True, fillcolor=255)
        text = pytesseract.image_to_string(candidate, lang="rus+eng", config=config)
        score = text_score(text)
        if score > best_score:
            best_text = text
            best_score = score

    return best_text


def normalize_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def ocr_file(file_path: Path) -> str:
    _, tessdata = configure_tesseract()
    with tempfile.TemporaryDirectory(prefix="passport_contract_") as temp:
        temp_dir = Path(temp)
        if file_path.suffix.lower() == ".pdf":
            image_paths = render_pdf(file_path, temp_dir)
        else:
            image_paths = [file_path]

        texts = [ocr_image(path, tessdata) for path in image_paths]
        return normalize_text("\n\n--- СТРАНИЦА ---\n\n".join(texts))


def normalize_date(value: str) -> str:
    cleaned = value.strip().replace("/", ".").replace("-", ".")
    match = re.fullmatch(r"(\d{2})\.(\d{2})\.(\d{2}|\d{4})", cleaned)
    if not match:
        return cleaned
    day, month, year_text = match.groups()
    if len(year_text) == 2:
        year = int(year_text)
        year_text = str(1900 + year if year >= 30 else 2000 + year)
    return f"{day}.{month}.{year_text}"


def find_date_near(text: str, keywords: tuple[str, ...], window: int = 260) -> str:
    lower = text.lower()
    for keyword in keywords:
        index = lower.find(keyword.lower())
        if index >= 0:
            fragment = text[max(0, index - 30): index + window]
            match = re.search(r"\b(\d{2}[.\-/]\d{2}[.\-/]\d{2,4})\b", fragment)
            if match:
                return normalize_date(match.group(1))
    return ""


def birth_date_from_mrz(text: str) -> str:
    compact_lines = [re.sub(r"\s+", "", line.upper()) for line in text.splitlines()]
    for line in compact_lines:
        if "RUS" not in line:
            continue
        match = re.search(r"RUS(\d{6})\d?[MFМЖ]", line)
        if not match:
            continue
        yy, mm, dd = match.group(1)[:2], match.group(1)[2:4], match.group(1)[4:6]
        current_yy = date.today().year % 100
        century = 1900 if int(yy) > current_yy else 2000
        return f"{dd}.{mm}.{century + int(yy)}"
    return ""


def mrz_to_cyrillic(token: str) -> str:
    token = re.sub(r"[^A-Z0-9]", "", token.upper())
    # Частая OCR-ошибка: цифра 9 (Я) читается как 8 (Ю).
    if token.endswith("I8") and len(token) >= 5:
        token = token[:-1] + "9"
    mapping = {
        "A": "А", "B": "Б", "V": "В", "G": "Г", "D": "Д",
        "E": "Е", "2": "Ё", "J": "Ж", "Z": "З", "I": "И",
        "Q": "Й", "K": "К", "L": "Л", "M": "М", "N": "Н",
        "O": "О", "P": "П", "R": "Р", "S": "С", "T": "Т",
        "U": "У", "F": "Ф", "H": "Х", "C": "Ц", "3": "Ч",
        "W": "Ш", "X": "Щ", "4": "Ъ", "Y": "Ы", "6": "Ь",
        "7": "Э", "8": "Ю", "9": "Я",
    }
    return "".join(mapping.get(char, "") for char in token).capitalize()


def fio_from_mrz(text: str) -> str:
    for raw_line in text.splitlines():
        line = re.sub(r"\s+", "", raw_line.upper())
        marker = line.find("PNRUS")
        if marker < 0:
            continue
        names = line[marker + 5:]
        if "<<" not in names:
            continue
        surname_raw, rest = names.split("<<", 1)
        other = [part for part in rest.split("<") if part]
        if len(other) < 2:
            continue
        surname = mrz_to_cyrillic(surname_raw)
        given = mrz_to_cyrillic(other[0])
        patronymic = mrz_to_cyrillic(other[1])
        if surname and given and patronymic:
            return f"{surname} {given} {patronymic}"
    return ""


def extract_fio(lines: list[str], birth_date: str, full_text: str) -> str:
    end_index = len(lines)
    if birth_date:
        for index, line in enumerate(lines):
            if birth_date in line:
                end_index = index
                break

    start_index = max(0, end_index - 18)
    area = lines[start_index:end_index + 2]
    candidates: list[str] = []

    for line in area:
        for word in re.findall(r"\b[А-ЯЁ]{3,}(?:-[А-ЯЁ]{2,})?\b", line.upper()):
            if word in NAME_STOPWORDS:
                continue
            if any(stop in word for stop in ("ФЕДЕРАЦ", "ПОДРАЗДЕЛ", "ПЕТЕРБУР")):
                continue
            candidates.append(word)

    unique: list[str] = []
    for word in candidates:
        if word not in unique:
            unique.append(word)

    mrz_fio = fio_from_mrz(full_text)
    if mrz_fio:
        return mrz_fio

    if len(unique) >= 3:
        selected = unique[-3:]
        return " ".join(word.capitalize() for word in selected)

    if unique:
        return " ".join(word.capitalize() for word in unique[-3:])
    return ""


def extract_issuer(lines: list[str], issue_date: str) -> str:
    issue_index = -1
    for index, line in enumerate(lines):
        if issue_date and issue_date in line:
            issue_index = index
            break

    search_area = lines[: issue_index + 1] if issue_index >= 0 else lines[:15]
    issuer_lines: list[str] = []
    for line in search_area:
        clean = re.sub(r"[^А-ЯЁа-яё0-9 .\-№]", " ", line)
        clean = re.sub(r"\s+", " ", clean).strip(" .-")
        upper = clean.upper()
        if len(re.findall(r"[А-ЯЁ]", upper)) < 4:
            continue
        if "РОССИЙСКАЯ ФЕДЕРАЦИЯ" in upper:
            continue
        if any(word in upper for word in ISSUER_WORDS):
            issuer_lines.append(clean)

    if issuer_lines:
        issuer = " ".join(issuer_lines)
        issuer = re.sub(r"\b\d{2}[.\-/]\d{2}[.\-/]\d{2,4}\b", "", issuer)
        issuer = re.sub(r"\b\d{3}-\d{3}\b", "", issuer)
        first_authority = re.search(
            r"\b(ОТДЕЛА|ОТДЕЛОМ|ОТДЕЛЕНИЕМ|ГУ\s+МВД|УФМС|МВД)\b",
            issuer.upper(),
        )
        if first_authority:
            issuer = issuer[first_authority.start():]
        return normalize_text(issuer)
    return ""


def extract_address(lines: list[str]) -> str:
    start = -1
    for index, line in enumerate(lines):
        upper = line.upper()
        if any(marker in upper for marker in (
            "ЗАРЕГИСТРИРОВАН", "МЕСТО ЖИТЕЛЬСТВА", "АДРЕС РЕГИСТРАЦИИ",
        )):
            start = index
            break

    if start < 0:
        return ""

    collected: list[str] = []
    for line in lines[start:start + 10]:
        clean = re.sub(r"\s+", " ", line).strip()
        if clean:
            collected.append(clean)

    result = " ".join(collected)
    result = re.sub(
        r"^(ЗАРЕГИСТРИРОВАН(?:А)?|МЕСТО ЖИТЕЛЬСТВА|АДРЕС РЕГИСТРАЦИИ)\s*[:\-]?\s*",
        "",
        result,
        flags=re.IGNORECASE,
    )
    return normalize_text(result)


def extract_passport_data(text: str) -> dict[str, str]:
    values = {key: "" for key, _ in FIELDS}
    lines = [normalize_text(line) for line in text.splitlines() if normalize_text(line)]
    flat = re.sub(r"\s+", " ", text)

    passport_candidates = re.findall(r"(?<!\d)(\d{4})\s*(\d{6})(?!\d)", flat)
    if passport_candidates:
        series, number = passport_candidates[-1]
        values["passport_series_number"] = f"{series} {number}"

    department = re.search(r"\b\d{3}\s*[-–]\s*\d{3}\b", flat)
    if department:
        values["department_code"] = re.sub(r"\s+", "", department.group(0)).replace("–", "-")

    values["issue_date"] = find_date_near(
        text, ("дата выдачи", "выдан", "выдано", "уфмс", "мвд")
    )
    values["birth_date"] = find_date_near(
        text, ("дата рождения", "родился", "родилась", "жен.", "муж.")
    ) or birth_date_from_mrz(text)

    values["fio"] = extract_fio(lines, values["birth_date"], text)
    values["issued_by"] = extract_issuer(lines, values["issue_date"])
    values["registration_address"] = extract_address(lines)
    return values


def replace_placeholder_in_paragraph(paragraph, placeholder: str, value: str) -> None:
    while True:
        runs = paragraph.runs
        full = "".join(run.text for run in runs)
        start = full.find(placeholder)
        if start < 0:
            return
        end = start + len(placeholder)

        positions = []
        cursor = 0
        for index, run in enumerate(runs):
            next_cursor = cursor + len(run.text)
            positions.append((index, cursor, next_cursor))
            cursor = next_cursor

        start_run = end_run = None
        start_offset = end_offset = 0
        for index, left, right in positions:
            if start_run is None and left <= start < right:
                start_run = index
                start_offset = start - left
            if left < end <= right:
                end_run = index
                end_offset = end - left
                break

        if start_run is None:
            if not runs:
                paragraph.add_run(value)
            else:
                runs[0].text = full.replace(placeholder, value, 1)
                for run in runs[1:]:
                    run.text = ""
            continue

        if end_run is None:
            end_run = len(runs) - 1
            end_offset = len(runs[end_run].text)

        prefix = runs[start_run].text[:start_offset]
        suffix = runs[end_run].text[end_offset:]
        runs[start_run].text = prefix + value + suffix
        for index in range(start_run + 1, end_run + 1):
            runs[index].text = ""


def replace_in_document(document: Document, values: dict[str, str]) -> None:
    placeholders = {f"{{{{{key}}}}}": value for key, value in values.items()}

    def process_paragraphs(paragraphs):
        for paragraph in paragraphs:
            for placeholder, value in placeholders.items():
                replace_placeholder_in_paragraph(paragraph, placeholder, value)

    process_paragraphs(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    for section in document.sections:
        for part in (section.header, section.footer):
            process_paragraphs(part.paragraphs)
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        process_paragraphs(cell.paragraphs)


def fill_template(template_path: Path, output_path: Path, values: dict[str, str]) -> None:
    document = Document(str(template_path))
    replace_in_document(document, values)
    document.save(str(output_path))


class ScrollableForm(ttk.Frame):
    def __init__(self, parent):
        super().__init__(parent)
        canvas = tk.Canvas(self, highlightthickness=0)
        scrollbar = ttk.Scrollbar(self, orient="vertical", command=canvas.yview)
        self.inner = ttk.Frame(canvas)
        self.inner.bind(
            "<Configure>",
            lambda _event: canvas.configure(scrollregion=canvas.bbox("all")),
        )
        canvas.create_window((0, 0), window=self.inner, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")


class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title(APP_TITLE)
        self.geometry("900x820")
        self.minsize(780, 680)

        self.passport_path = tk.StringVar()
        default_template = app_dir() / "template_gph.docx"
        self.template_path = tk.StringVar(
            value=str(default_template) if default_template.exists() else ""
        )
        desktop = Path.home() / "Desktop"
        self.output_dir = tk.StringVar(value=str(desktop if desktop.exists() else Path.home()))
        self.status = tk.StringVar(value="Готово к работе")
        self.raw_ocr = ""
        self.vars = {key: tk.StringVar() for key, _ in FIELDS}
        self.vars["contract_date"].set(date.today().strftime("%d.%m.%Y"))
        self._build_ui()

    def _build_ui(self):
        ttk.Label(self, text=APP_TITLE, font=("Segoe UI", 17, "bold")).pack(pady=(12, 5))
        ttk.Label(
            self,
            text="Все файлы обрабатываются локально. Перед созданием договора обязательно проверь поля.",
        ).pack(pady=(0, 10))

        files = ttk.LabelFrame(self, text="Файлы")
        files.pack(fill="x", padx=12, pady=5)
        self._file_row(files, 0, "Паспорт PDF/фото", self.passport_path, self.choose_passport)
        self._file_row(files, 1, "Шаблон DOCX", self.template_path, self.choose_template)
        self._file_row(files, 2, "Папка сохранения", self.output_dir, self.choose_output_dir)

        buttons = ttk.Frame(self)
        buttons.pack(fill="x", padx=12, pady=8)
        self.recognize_button = ttk.Button(
            buttons, text="1. Распознать паспорт", command=self.recognize
        )
        self.recognize_button.pack(side="left", padx=(0, 8))
        ttk.Button(buttons, text="Показать текст OCR", command=self.show_ocr).pack(
            side="left", padx=(0, 8)
        )
        ttk.Button(buttons, text="2. Создать договор", command=self.create_contract).pack(
            side="left"
        )

        form_box = ttk.LabelFrame(self, text="Проверь и дополни данные")
        form_box.pack(fill="both", expand=True, padx=12, pady=5)
        scroll = ScrollableForm(form_box)
        scroll.pack(fill="both", expand=True, padx=6, pady=6)

        for row, (key, label) in enumerate(FIELDS):
            ttk.Label(scroll.inner, text=label, width=27).grid(
                row=row, column=0, sticky="w", padx=8, pady=4
            )
            ttk.Entry(scroll.inner, textvariable=self.vars[key], width=74).grid(
                row=row, column=1, sticky="ew", padx=8, pady=4
            )
        scroll.inner.grid_columnconfigure(1, weight=1)

        ttk.Label(
            self,
            textvariable=self.status,
            relief="sunken",
            anchor="w",
        ).pack(fill="x", side="bottom")

    def _file_row(self, parent, row, label, variable, command):
        ttk.Label(parent, text=label, width=21).grid(
            row=row, column=0, sticky="w", padx=8, pady=5
        )
        ttk.Entry(parent, textvariable=variable).grid(
            row=row, column=1, sticky="ew", padx=5, pady=5
        )
        ttk.Button(parent, text="Выбрать", command=command).grid(
            row=row, column=2, padx=8, pady=5
        )
        parent.grid_columnconfigure(1, weight=1)

    def choose_passport(self):
        path = filedialog.askopenfilename(
            filetypes=[
                ("Паспорт PDF или фото", "*.pdf *.png *.jpg *.jpeg *.tif *.tiff"),
                ("Все файлы", "*.*"),
            ]
        )
        if path:
            self.passport_path.set(path)

    def choose_template(self):
        path = filedialog.askopenfilename(filetypes=[("Документ Word", "*.docx")])
        if path:
            self.template_path.set(path)

    def choose_output_dir(self):
        path = filedialog.askdirectory()
        if path:
            self.output_dir.set(path)

    def set_busy(self, busy: bool, text: str):
        self.status.set(text)
        self.recognize_button.configure(state="disabled" if busy else "normal")
        self.configure(cursor="watch" if busy else "")

    def recognize(self):
        path = Path(self.passport_path.get())
        if not path.exists():
            messagebox.showerror("Нет файла", "Выбери PDF или фотографию паспорта.")
            return

        self.set_busy(True, "Распознаю паспорт — это может занять 20–90 секунд…")
        threading.Thread(target=self._recognize_worker, args=(path,), daemon=True).start()

    def _recognize_worker(self, path: Path):
        try:
            text = ocr_file(path)
            data = extract_passport_data(text)
            self.after(0, lambda: self._recognize_done(text, data))
        except Exception as error:
            error_text = f"{type(error).__name__}: {error!r}"
            self.after(0, lambda message=error_text: self._recognize_failed(message))

    def _recognize_done(self, text: str, data: dict[str, str]):
        self.raw_ocr = text
        for key, value in data.items():
            if value:
                self.vars[key].set(value)
        self.set_busy(False, "Распознавание завершено. Проверь каждое поле.")
        messagebox.showinfo(
            "Паспорт распознан",
            "Проверь ФИО, номер паспорта, даты, орган выдачи и адрес.\n"
            "Если адреса нет на загруженной странице, поле останется пустым.",
        )

    def _recognize_failed(self, error_message: str):
        self.set_busy(False, "Ошибка распознавания")
        messagebox.showerror(
            "Ошибка OCR",
            error_message or "Неизвестная ошибка OCR. Сообщение отсутствует.",
        )

    def show_ocr(self):
        if not self.raw_ocr:
            messagebox.showinfo("Текст OCR", "Сначала распознай паспорт.")
            return
        window = tk.Toplevel(self)
        window.title("Распознанный текст")
        window.geometry("800x600")
        text = tk.Text(window, wrap="word")
        text.insert("1.0", self.raw_ocr)
        text.pack(fill="both", expand=True)
        text.configure(state="disabled")

    def create_contract(self):
        template = Path(self.template_path.get())
        output_dir = Path(self.output_dir.get())

        if not template.exists() or template.suffix.lower() != ".docx":
            messagebox.showerror("Нет шаблона", "Выбери шаблон договора в формате DOCX.")
            return

        values = {key: variable.get().strip() for key, variable in self.vars.items()}
        if not values["fio"]:
            if not messagebox.askyesno(
                "ФИО не заполнено",
                "Поле ФИО пустое. Всё равно создать договор?",
            ):
                return

        output_dir.mkdir(parents=True, exist_ok=True)
        safe_name = re.sub(
            r"[^А-Яа-яЁёA-Za-z0-9_-]+",
            "_",
            values["fio"] or "Исполнитель",
        ).strip("_")
        output = output_dir / f"Договор_ГПХ_{safe_name}.docx"

        try:
            fill_template(template, output, values)
            self.status.set(f"Создан файл: {output}")
            messagebox.showinfo("Договор создан", f"Готовый файл:\n{output}")
            if os.name == "nt":
                os.startfile(output_dir)
        except Exception as error:
            messagebox.showerror("Ошибка заполнения", str(error))


if __name__ == "__main__":
    App().mainloop()
